
from openai import OpenAI
from src.config.config import OPENAI_API_KEY
from fastapi import status
from docx import Document
from docx.shared import Pt
import json
import datetime

class Analysis_ServiceOld():

    def __init__(self) -> None:

        self.openai_key=OPENAI_API_KEY
        self.client = OpenAI(api_key=self.openai_key)
        self.model = "gpt-4o"
        self.context_path="assets\Context.txt"


    def read_text_file(self, file_path):
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                # Reading the contents of the file
                content = file.read()
                return content
        except FileNotFoundError:
            print(f"The file at {file_path} was not found.")
        except Exception as e:
            print(f"An error occurred: {e}")


    def convert_json_to_docx(self, json_data, file_name):
        """
        Converts the given JSON data into a formatted Word document.
        
        Parameters:
        json_data (dict): The JSON data to convert.
        file_name (str): The name of the file to save.
        """
        # Create a new Document
        doc = Document()
        
        # Add title
        doc.add_heading('Medical Report', 0)
        
        # Add Pathology Report Summary
        doc.add_heading('Pathology Report Summary', level=1)
        pathology_report = json_data.get("Pathology Report Summary", {})
        for key, value in pathology_report.items():
            # Handle nested dictionaries like 'Histologic Score' and 'Receptor Status'
            if isinstance(value, dict):
                doc.add_paragraph(f"{key}:")
                for sub_key, sub_value in value.items():
                    doc.add_paragraph(f"    {sub_key}: {sub_value}", style='List Bullet')
            else:
                doc.add_paragraph(f"{key}: {value}")

        # Add Pedigree Analysis
        doc.add_heading('Pedigree Analysis', level=1)
        pedigree_analysis = json_data.get("Pedigree Analysis", {})
        
        for generation, details in pedigree_analysis.items():
            doc.add_paragraph(f"{generation}:")
            
            # Check if the details are a list (multiple family members)
            if isinstance(details, list):
                for member in details:
                    for key, value in member.items():
                        doc.add_paragraph(f"    {key}: {value}", style='List Bullet')
            # Otherwise, handle as a dictionary for individual family members
            elif isinstance(details, dict):
                for key, value in details.items():
                    if isinstance(value, dict):  # Nested dictionary for Patient details
                        doc.add_paragraph(f"    {key}:")
                        for sub_key, sub_value in value.items():
                            doc.add_paragraph(f"        {sub_key}: {sub_value}", style='List Bullet')
                    else:
                        _values = ", ".join(value)
                        doc.add_paragraph(f"    {key}: {_values}", style='List Bullet')

        # Add NCCN Testing Criteria Assessment
        doc.add_heading('NCCN Testing Criteria Assessment', level=1)
        nccn_assessment = json_data.get("NCCN Testing Criteria Assessment", {})
        for key, value in nccn_assessment.items():

            _values = ", ".join(value)
            doc.add_paragraph(f"{key}: {_values}")

        # Add Conclusion
        doc.add_heading('Conclusion', level=1)
        conclusion = json_data.get("Conclusion", {})

        if isinstance(conclusion, dict): 
            for key, value in conclusion.items():

                if isinstance(value, list):
                    _values = ", ".join(value)
                    doc.add_paragraph(_values)
                else:
                    doc.add_paragraph(value)
        else : 
            doc.add_paragraph(f"{conclusion}")
        # Save the document
        doc.save(file_name)
        return file_name


    def convert_to_json_object(self, json_string: str):
        
        try:
            # Find the index of the first occurrence of '{' and the last occurrence of '}'
            start_index = json_string.find("{")
            end_index = json_string.rfind("}")
            
            # If '{' or '}' is not found, raise an error
            if start_index == -1 or end_index == -1:
                print("Error: No valid JSON object found in the string.")
                return None
            
            # Remove all text before the first '{' and after the last '}'
            cleaned_string = json_string[start_index:end_index+1].strip()
            
            
            # Convert the cleaned string to a Python dictionary (JSON object)
            json_object = json.loads(cleaned_string)
            return json_object
        
        except json.JSONDecodeError as e:
            print(f"Error decoding JSON: {e}")
            return None


    def convert_path_to_url(self, file_path: str) -> str:
        """
        Convert a relative file path to a full URL for downloading through FastAPI's static files.

        Args:
            file_path (str): The relative path to the file.

        Returns:
            str: The full URL pointing to the static file.
        """
        base_url = "http://127.0.0.1:8000/static/"
        
        # Replace 'assets' with 'static' to create the proper URL path
        url_path = file_path.replace("assets/", "")
        
        # Combine the base URL with the cleaned file path
        full_url = f"{base_url}{url_path}"
        
        return full_url


    def read_json_file(self, file_path):
        """
        Reads a JSON file and returns its content.
        
        Args:
        - file_path (str): The path to the JSON file.
        
        Returns:
        - dict: Parsed JSON content.
        """

        with open(file_path, 'r') as file:
                data = json.load(file)
        return data
    
        try:
            with open(file_path, 'r') as file:
                data = json.load(file)
            return data
        except FileNotFoundError:
            print(f"Error: File not found at {file_path}.")
            return None
        except json.JSONDecodeError:
            print("Error: Failed to parse JSON file.")
            return None
        except Exception as e:
            print(f"An error occurred while reading the file: {e}")
            return None


    def create_prompt(self, frame, docs):
        """
        Create a prompt for analysis based on the given image, document, and an example from a JSON file.
        
        Args:
        - frame (str): Base64 encoded image.
        - docs (str): The document text for analysis.
        - example_file (str): Path to the example JSON file.
        
        Returns:
        - dict: A structured prompt including the image, document, rules, and example.
        """
        # Read the example JSON content from the file
        example_data = self.read_json_file("src/services/example.json")
        evaluation_rules = self.read_text_file(self.context_path)
        example_context = json.dumps(example_data, indent=4) if example_data else "No example available."
        
        # Prepare the prompt content with the image, document, and example for analysis
        prompt_context = (
            f"Context: Perform an analysis based on the provided document and image, using the guidelines and example provided below for structure and content.\n\n"
            f"Document: {docs}\n\n"
            f"Rules:\n"
            f"Evaluation Rules: \n{evaluation_rules}\n\n"
            f"Analysis Instructions: Please conduct a comprehensive analysis of the document and accompanying image. Format the results as a JSON object with the following sections, detailed as specified:\n"
            f"1. 'Pathology Report Summary': Include anonymized patient identifiers and the main findings from the pathology report, such as diagnosis, tumor characteristics (size, grade, stage), and any suggested treatment approaches.\n"
            f"2. 'Pedigree Analysis': Summarize the pedigree analysis by detailing the family's genetic health history. This should include both a visual representation (as a link to an image or a detailed description) and a textual analysis identifying significant familial links and any noted hereditary conditions.\n\n"
            f"3. 'NCCN Testing Criteria Assessment': Evaluate and determine the patient's eligibility for NCCN testing based on the listed criteria in the document. Provide a list of these specifying which are met.\n"
            f"4. 'Conclusion': should mentioned clearly if the patient is eligible or not for the NCCN Criteria then add all details you want.\n"
            f"Ensure each section is clearly labeled and organized to effectively present all relevant details. The response should adhere to JSON format standards, with appropriate data types and hierarchical structuring to clearly convey the information.\n"
            f"Example:\n{example_context}\n\n"
        )

        # Return the structured prompt
        return {
            "role": "user",
            "content": [
                prompt_context,
                {"type": "image_url", "image_url": {"url": f"data:image/jpeg;base64,{frame}"}}
            ],
        }
    
    
    def generate_file_name(self, base_name="medical_report", extension=".docx"):
        """
        Generates a unique file name with a timestamp.
        
        Parameters:
        base_name (str): The base name of the file. Default is 'medical_report'.
        extension (str): The file extension. Default is '.docx'.
        
        Returns:
        str: A file name in the format 'base_name_YYYYMMDD_HHMMSS.extension'.
        """
        timestamp = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
        file_name = f"{base_name}_{timestamp}{extension}"
        return file_name

    def generate_ai_report(self, frame, docs):

        prompt_context=self.create_prompt(frame, docs)
        # Parameters for API call
        params = {
            "model": self.model,
            "messages": [prompt_context],
        }

        # Make the API call
        result = self.client.chat.completions.create(**params)
        content= result.choices[0].message.content
        content = self.convert_to_json_object(content)
        file_name = self.generate_file_name()
        docx_path = self.convert_json_to_docx(content, f"assets/ai_reports/{file_name}")
        docx_path = {
            "docx_path": self.convert_path_to_url(docx_path),
            "content":content
            }
        return docx_path, status.HTTP_200_OK
    

analysis_service = Analysis_ServiceOld()