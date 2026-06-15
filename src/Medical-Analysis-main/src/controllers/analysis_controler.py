from io import BytesIO
import base64
from ..config.config import DB_HOST, DB_USER, DB_PASSWORD, DB_NAME
import logging
from logging.handlers import RotatingFileHandler

from fastapi.encoders import jsonable_encoder
from fastapi import APIRouter, File, UploadFile, status, Request, HTTPException, Form
from fastapi.responses import HTMLResponse, JSONResponse
from fastapi.templating import Jinja2Templates
from src.services.analysis_service import analysis_service
from typing import Optional
from docx import Document
import mysql.connector
from PIL import Image
import PyPDF2
from pydantic import BaseModel

class AnalysisRequest(BaseModel):
    userId: Optional[int]
    useUploadedFile: Optional[bool]



# ✅ Setup logger properly
logger = logging.getLogger("medical_report_logger")
logger.setLevel(logging.DEBUG)

# Rotating File Handler
rfh = RotatingFileHandler('app.log', maxBytes=5*1024*1024, backupCount=3)
rfh.setLevel(logging.DEBUG)
formatter = logging.Formatter('%(asctime)s - %(name)s - %(levelname)s - %(message)s')
rfh.setFormatter(formatter)
logger.addHandler(rfh)

# Console handler
ch = logging.StreamHandler()
ch.setLevel(logging.DEBUG)
ch.setFormatter(formatter)
logger.addHandler(ch)


MedicalRouter = APIRouter()
templates = Jinja2Templates(directory="templates")

# Function to read and extract text from the docx file
def read_docx_file(file_content): 
    try:
        # Use BytesIO to read the content into a docx Document object
        docs = Document(BytesIO(file_content))
        # Extract and return the text content from the .docx file
        doc_text = "\n".join([para.text for para in docs.paragraphs])
        return doc_text
    except Exception as e:
        logger.error(f"Error reading .docx file: {str(e)}")  # Log the error
        return f"Error reading .docx file: {str(e)}"

# @MedicalRouter.post("/analysis/")
# async def analysis(userId: Optional[int] = Form(None), useUploadedFile: Optional[bool] = Form(None)):
#     logger.info(f"Starting analysis for userId: {userId}")
    
@MedicalRouter.post("/analysis/")
async def analysis(request: AnalysisRequest):
    logger.info(f"Starting analysis for userId: {request.userId}")
    try:
        userId = request.userId
        useUploadedFile = request.useUploadedFile
        db_config = {
        "host": DB_HOST,
        "user": DB_USER,
        "password": DB_PASSWORD,
        "database": DB_NAME
        }
        service = analysis_service(db_config)
        db_data = service.get_user_data(userId)
        logger.info(f"Data retrieved for userId: {userId}")

        if not db_data:
            logger.warning("User data not found.")
            raise HTTPException(status_code=404, detail="No data found for the provided userId.")

        pedigree = db_data.get("pedigree")
        clinical = db_data.get("clinical")

        if not pedigree:
            logger.warning("No pedigree data.")
            raise HTTPException(status_code=404, detail="No pedigree data found.")

        img_data = pedigree.get("img")

        if not clinical or not clinical.get("histologyReportContent"):
            logger.warning("No clinical history.")
            raise HTTPException(status_code=404, detail="No clinical history report found.")

        frame = None
        docs = None
        filetype = pedigree.get("filetype", "").lower()

        if useUploadedFile is True:
            # Use image from `pedigree.img`

            if img_data:
                if filetype == "application/pdf":
                    try:
                        file_content = base64.b64decode(img_data)
                        pdf_reader = PyPDF2.PdfReader(BytesIO(file_content))
                        text = ""
                        for page in pdf_reader.pages:
                            text += page.extract_text()
                        frame = text
                        logger.info("PDF text extracted.")
                    except Exception as e:
                        logger.error(f"Error processing PDF: {e}")
                        raise HTTPException(status_code=500, detail=f"Error processing PDF: {e}")

                elif filetype.startswith("image/"):
                    frame = base64.b64encode(img_data).decode("utf-8")
                    logger.info("Image from pedigree encoded as base64.")

                elif filetype == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
                    try:
                        file_content = base64.b64decode(img_data)
                        doc = Document(BytesIO(file_content))
                        frame = "\n".join([p.text for p in doc.paragraphs])
                        logger.info("DOCX text extracted.")
                    except Exception as e:
                        logger.error(f"Error processing DOCX: {e}")
                        raise HTTPException(status_code=500, detail=f"Error processing DOCX: {e}")
                else:
                    logger.warning(f"Unsupported file type in 'img': {filetype}")

        elif useUploadedFile is False:
            # Use JSON text from `pedigree.csv`
            csv_data = pedigree.get("csv")
            if csv_data:
                frame = csv_data
                logger.info("CSV from pedigree loaded as JSON string.")
                filetype = 'json'
            else:
                if img_data :
                    if filetype == "application/pdf":
                        try:
                            file_content = base64.b64decode(img_data)
                            pdf_reader = PyPDF2.PdfReader(BytesIO(file_content))
                            text = ""
                            for page in pdf_reader.pages:
                                text += page.extract_text()
                            frame = text
                            logger.info("PDF text extracted.")
                        except Exception as e:
                            logger.error(f"Error processing PDF: {e}")
                            raise HTTPException(status_code=500, detail=f"Error processing PDF: {e}")

                    elif filetype.startswith("image/"):
                        frame = base64.b64encode(img_data).decode("utf-8")
                        logger.info("Image from pedigree encoded as base64.")

                    elif filetype == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
                        try:
                            file_content = base64.b64decode(img_data)
                            doc = Document(BytesIO(file_content))
                            frame = "\n".join([p.text for p in doc.paragraphs])
                            logger.info("DOCX text extracted.")
                        except Exception as e:
                            logger.error(f"Error processing DOCX: {e}")
                            raise HTTPException(status_code=500, detail=f"Error processing DOCX: {e}")
                    else:
                        logger.warning("No CSV data found in pedigree.")

        else:
            img_data = pedigree.get("img")
            if img_data:
                frame = base64.b64encode(img_data).decode("utf-8")
                logger.info("Image from pedigree used as fallback and encoded as base64.")
                filetype = "image"
            else:
                csv_data = pedigree.get("csv")
                if csv_data:
                    frame = csv_data
                    logger.info("CSV from pedigree used as fallback and loaded as JSON string.")
                    filetype = "json"
                else:
                    logger.warning("No image or CSV data found in pedigree.")
                    frame = None
                    filetype = None

        try:
            docx_obj = Document(BytesIO(clinical["histologyReportContent"]))
            docs = "\n".join([p.text for p in docx_obj.paragraphs])
            logger.info("DOCX content parsed from clinical report.")
        except Exception as e:
            logger.error(f"Failed to parse DOCX content: {str(e)}")
            raise HTTPException(status_code=500, detail="Failed to parse clinical DOCX file.")

        logger.info("Preparing to generate AI report with the following parameters:")
        logger.info(f"User ID: {userId}")
        logger.info(f"Use Uploaded File: {useUploadedFile}")
        logger.info(f"Filetype: {filetype}")
        logger.info(f"Docs: {docs}")

        result = service.generate_ai_report(
            frame=frame,
            docs=docs,
            userId=userId,
            filetype=filetype,
            userUploadedFile=useUploadedFile
        )

        logger.info("AI report generated successfully.")
        return JSONResponse(content=result, status_code=result.get("status_code", 200))

    except mysql.connector.Error as e:
        logger.error(f"Database error: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Database error: {str(e)}")
    except Exception as e:
        error_message = str(e).strip() or "Unknown internal error"
        logger.error(f"Unexpected error: {error_message}")
        raise HTTPException(status_code=500, detail=f"{error_message}")

    
# @MedicalRouter.post("/analysis/")
# async def analysis(userId: int = Form(...), useUploadedFile: bool = Form(...)):
#     logger.info(f"Starting analysis for userId: {userId}")

#     try:
#         service = analysis_service()
#         db_data = service.get_user_data(userId)
#         logger.info(f"service.get_user_data userId: {userId}")
#         if not db_data:
#             logger.warning("User data not found.")
#             raise HTTPException(status_code=404, detail="No data found for the provided userId.")

#         # logger.info(f"Start gettings pedigree")

#         pedigree = db_data.get("pedigree")
#         # logger.info(f"==End  gettings pedigree")

#         # logger.info(f"Start gettings clinical")

#         clinical = db_data.get("clinical")
#         # logger.info(f"==End  gettings clinical")


#         if not pedigree:
#             logger.warning("No pedigree data.")
#             raise HTTPException(status_code=404, detail="No pedigree data found.")

#         frame = None
#         docs = None
#         filetype = pedigree.get("filetype", "").lower()
#         mime_type = None

#         if useUploadedFile == True:
        
#             # Ensure 'filetype' is available
#             if pedigree and "filetype" in pedigree:
#                 ext_to_mime = {
#                     "jpg": "image/jpeg",
#                     "jpeg": "image/jpeg",
#                     "png": "image/png",
#                     "gif": "image/gif",
#                     "webp": "image/webp",
#                     "pdf": "application/pdf"
#                 }
#                 extension = ext_to_mime.get(filetype, "bin")  # default to 'bin' for unknown types
#                 mime_type = filetype if filetype in mime_to_ext else "application/octet-stream"
#         else:
#             extension = "json"
#             mime_type = "application/json"

#         # Handle image
#         if not clinical or not clinical.get("histologyReportContent"):
#             logger.warning("No clinical history.")
#             raise HTTPException(status_code=404, detail="No clinical history report found.")
        
#         if extension in ext_to_mime and mime_type.startswith("image") and pedigree.get("img"):
#             frame = base64.b64encode(clinical["histologyReportContent"]).decode("utf-8")
#             logger.info(f"Image encoded to base64 with MIME type {clinical["histologyReportType"] }.")

#         # Handle json
#         elif extension == "json" and pedigree.get("csv"):
#             docs = pedigree["csv"]
#             docs_blob = clinical["histologyReportContent"]
#         try:
#             docx_obj = Document(BytesIO(docs_blob))
#             docs = "\n".join([p.text for p in docx_obj.paragraphs])
#             logger.info("DOCX file parsed.")
#         except Exception as e:
#             logger.error(f"Failed to parse DOCX: {str(e)}")
#             logger.info("json content loaded.")

#         else:
#             logger.warning("No valid filetype found in pedigree data.")


    
        
        

#         # filetype=pedigree.get("filetype") 
#         # logger.info(f"File type: {filetype}")
#         result  = service.generate_ai_report(
#             frame=frame,
#             docs=docs,
#             userId=userId,
#             filetype=filetype,
#             userUploadedFile=useUploadedFile
#         )

#         payload = result
#         status_code = result.get("status_code", 500)
#         logger.info("AI report generated.")
        
#         return JSONResponse(content=payload, status_code=status_code)

#     except mysql.connector.Error as e:
#         logger.error(f"DB error: {str(e)}")
#         raise HTTPException(status_code=500, detail="Database error.")
#     except Exception as e:
#         logger.error(f"Unexpected error: {str(e)}")
#         raise HTTPException(status_code=500, detail="Internal server error.")

# @MedicalRouter.post("/analysis/")
# async def analysis(userId: int = Form(...)):
#     db = None
#     cursor = None
#     try:
#         logger.info(f"Starting analysis for userId: {userId}")

#         # Connect to MySQL database
#         service = analysis_service()
#         db = service.get_db()
#         cursor = db.cursor(dictionary=True)

#         # Log database connection success
#         logger.info("Database connection established successfully.")

#         # Get data from family_pedigree
#         cursor.execute("""
#             SELECT img, csv, filetype
#             FROM lamadb.family_pedigree
#             WHERE userId = %s
#         """, (userId,))
#         pedigree_data = cursor.fetchone()

#         if not pedigree_data:
#             logger.warning(f"No pedigree data found for userId: {userId}")
#             raise HTTPException(
#                 status_code=status.HTTP_404_NOT_FOUND,
#                 detail="No pedigree data found for the provided userId."
#             )
        
#         logger.info(f"Pedigree data retrieved for userId: {userId}")
        
#         frame = None
#         csv_content = None
#         filetype = pedigree_data.get("filetype", "").lower()

#         if filetype == "image" and pedigree_data.get("img"):
#             frame = base64.b64encode(pedigree_data["img"]).decode("utf-8")
#             logger.info("Image data found and base64 encoded.")
#         elif filetype == "csv" and pedigree_data.get("csv"):
#             csv_content = pedigree_data["csv"]
#             logger.info("CSV data found.")
        
#         # Get histology report from clinical_histories
#         cursor.execute("""
#             SELECT histologyReportContent, histologyReportName, histologyReportType
#             FROM lamadb.clinical_histories
#             WHERE userId = %s
#         """, (userId,))
#         clinical_data = cursor.fetchone()

#         if not clinical_data or not clinical_data["histologyReportContent"]:
#             logger.warning(f"No clinical history report found for userId: {userId}")
#             raise HTTPException(
#                 status_code=status.HTTP_404_NOT_FOUND,
#                 detail="No clinical history report found for the provided userId."
#             )

#         # Extract .docx content from blob
#         docs = read_docx_file(clinical_data["histologyReportContent"])

#         # Log report generation
#         logger.info(f"Generating AI report for userId: {userId}.")

#         # Generate report
#         payload, status_code = service.generate_ai_report(
#             frame=frame,
#             docs=docs if docs else csv_content,
#             userId=userId
#         )

#         logger.info(f"Report successfully generated for userId: {userId}")
#         return JSONResponse(content=payload, status_code=status_code)

#     except mysql.connector.Error as e:
#         logger.error(f"Database error: {str(e)}")
#         raise HTTPException(
#             status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
#             detail=f"Database error: {str(e)}"
#         )
#     except Exception as e:
#         logger.error(f"Unexpected error: {str(e)}")
#         raise HTTPException(
#             status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
#             detail=f"Error: {str(e)}"
#         )
#     finally:
#         # Ensure the cursor and database connection are closed properly
#         if cursor:
#             cursor.close()
#             logger.info("Database cursor closed.")
#         if db:
#             db.close()
#             logger.info("Database connection closed.")
